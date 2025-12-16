#!/usr/bin/env python3
"""
Advanced Resume Screening System (Fixed & Optimized)
All critical fixes applied:
- Word boundary fix for C++, C#, .NET
- Pre-filtered fuzzy matching (performance)
- Weighted average for chunked semantics
- Empty variant handling
- PDF page limit off-by-one fix
- Additional optimizations
"""

import io
import os
import re
import time
import math
import unicodedata
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime

from concurrent.futures import ThreadPoolExecutor, as_completed

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

from pypdf import PdfReader
from docx import Document
from rapidfuzz import fuzz
from sentence_transformers import SentenceTransformer, util

# ------------------------- Configuration -------------------------

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_RESUMES = 50
MAX_PDF_PAGES = 100  # Process at most 100 pages per PDF
MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

# Enable chunked semantic similarity (more robust on long resumes)
USE_CHUNKED_SEMANTICS = True
SEMANTIC_CHUNK_CHAR_LEN = 1500
SEMANTIC_CHUNK_AGG = "weighted_top"  # "max", "avg", or "weighted_top" (recommended)
SEMANTIC_TOP_K = 3  # For weighted_top: use top 3 chunks

# Fuzzy matching optimization
FUZZY_PRE_FILTER = True  # Enable pre-filtering for performance
FUZZY_MAX_CHUNK_LEN = 1024  # Skip very long chunks

# Parallelism
MAX_PARSE_WORKERS = 4

# CORS: lock down in production
ALLOWED_ORIGINS = ["http://localhost:3000", "http://127.0.0.1:8000", "http://localhost:8000"]

# ------------------------- Knowledge Bases -------------------------

TECH_SYNONYMS = {
    # Programming Languages
    "javascript": ["js", "ecmascript", "es6", "es2015", "es2020"],
    "typescript": ["ts"],
    "python": ["py", "python3"],
    "java": [],
    "c++": ["cpp", "c plus plus", "cplusplus"],
    "c#": ["csharp", "c sharp"],
    "go": ["golang"],
    "php": [],
    "ruby": [],
    "rust": [],
    "kotlin": [],
    "swift": [],

    # Frameworks & Libraries
    "react": ["reactjs", "react.js"],
    "vue": ["vuejs", "vue.js"],
    "angular": ["angularjs"],
    "node.js": ["nodejs", "node"],
    "next.js": ["nextjs"],
    "express": ["expressjs", "express.js"],
    ".net": ["dotnet", "asp.net", "asp.net core", "aspnet"],
    "django": [],
    "flask": [],
    "spring": ["spring boot", "springboot"],

    # Databases
    "postgresql": ["postgres", "psql", "postgre", "postgre sql"],
    "mongodb": ["mongo"],
    "mysql": ["my sql"],
    "redis": [],
    "sqlite": [],
    "elasticsearch": ["elastic search"],

    # Cloud & DevOps
    "amazon web services": ["aws"],
    "google cloud": ["gcp", "google cloud platform"],
    "microsoft azure": ["azure"],
    "docker": ["containerization"],
    "kubernetes": ["k8s", "k-8-s"],
    "terraform": [],
    "ansible": [],
    "jenkins": [],
    "ci/cd": ["ci cd", "continuous integration", "continuous deployment"],

    # AI/ML
    "machine learning": ["ml"],
    "artificial intelligence": ["ai"],
    "natural language processing": ["nlp"],
    "deep learning": ["dl"],
    "tensorflow": ["tf"],
    "pytorch": ["torch"],

    # Other
    "application programming interface": ["api", "rest api", "restful"],
    "user interface": ["ui"],
    "user experience": ["ux"],
}

SENIORITY_INDICATORS = {
    "entry": ["intern", "junior", "entry level", "graduate", "fresh graduate", "0-2 years", "l1"],
    "mid": ["mid level", "intermediate", "2-5 years", "3-5 years", "experienced", "l2", "l3"],
    "senior": ["senior", "lead", "principal", "staff", "5+ years", "7+ years", "expert", "sr.", "l4", "l5"],
    "expert": ["architect", "director", "vp", "chief", "head of", "10+ years", "15+ years", "distinguished", "fellow"],
}

STOPWORDS = set("""
a an the and or but if while is are was were be been being to for on in of by 
with at from as that this these those it its into over under between about 
across via you your we our they their i me my he she will would could should 
may might must can have has had do does did having done being work experience
position role responsibilities company team project projects looking seeking
required preferred qualifications skills strong excellent good knowledge
understanding ability develop implement manage responsible deliver build design
""".split())

# Month aliases for date parsing
MONTHS = {
    "jan": 1, "january": 1, "janvier": 1,
    "feb": 2, "february": 2, "février": 2, "fevrier": 2,
    "mar": 3, "march": 3, "mars": 3,
    "apr": 4, "april": 4, "avril": 4,
    "may": 5, "mai": 5,
    "jun": 6, "june": 6, "juin": 6,
    "jul": 7, "july": 7, "juillet": 7,
    "aug": 8, "august": 8, "août": 8, "aout": 8,
    "sep": 9, "september": 9, "septembre": 9,
    "oct": 10, "october": 10, "octobre": 10,
    "nov": 11, "november": 11, "novembre": 11,
    "dec": 12, "december": 12, "décembre": 12, "decembre": 12,
}

# ------------------------- Data Models -------------------------

@dataclass
class ScoringWeights:
    semantic: float = 0.35
    keywords: float = 0.25
    must_haves: float = 0.20
    experience: float = 0.10
    education: float = 0.05
    bonus: float = 0.05

    def normalized(self) -> "ScoringWeights":
        s = self.semantic + self.keywords + self.must_haves + self.experience + self.education + self.bonus
        if s <= 0:
            return self
        return ScoringWeights(
            semantic=self.semantic / s,
            keywords=self.keywords / s,
            must_haves=self.must_haves / s,
            experience=self.experience / s,
            education=self.education / s,
            bonus=self.bonus / s,
        )

@dataclass
class ScreeningResult:
    candidate_name: str
    filename: str
    decision: str
    confidence: float
    total_score: float

    semantic_score: float
    keyword_coverage: float
    must_have_score: float
    experience_score: float
    education_score: float
    bonus_score: float

    years_experience: Optional[float]
    matched_keywords: List[str]
    missing_must_haves: List[str]
    matched_nice_to_haves: List[str]
    red_flags: List[str]
    strengths: List[str]

    reasoning: str

    def to_dict(self):
        return asdict(self)

# ------------------------- Utilities -------------------------

def normalize_text(s: str) -> str:
    """Normalize text for matching"""
    s = s or ""
    s = s.replace("\x00", " ")
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.lower()
    s = re.sub(r"[^\w\s\+#\./\-]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def read_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF with page limit fix"""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        for i, page in enumerate(reader.pages):
            if i >= MAX_PDF_PAGES:  # FIX: >= instead of > (was processing 101 pages)
                break
            try:
                text = page.extract_text()
                if text:
                    texts.append(text)
            except Exception:
                continue
        return "\n".join(texts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF parsing failed: {str(e)}")

def read_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including tables"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        
        # Paragraphs
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        
        # Tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells:
                    parts.append(" | ".join(cells))
        
        return "\n".join(parts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"DOCX parsing failed: {str(e)}")

def extract_text_from_upload(upload: UploadFile) -> Tuple[str, str]:
    """Extract text from uploaded file with validation"""
    filename = upload.filename or "file"
    ext = os.path.splitext(filename.lower())[1]

    content = upload.file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File {filename} exceeds 10MB limit")

    if ext == ".pdf":
        text = read_pdf(content)
    elif ext == ".docx":
        text = read_docx(content)
    elif ext in (".txt", ".md"):
        text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    upload.file.close()

    if not text.strip():
        raise HTTPException(status_code=400, detail=f"No text extracted from {filename}")

    return text, filename

# ------------------------- Keyword Extraction -------------------------

def expand_keyword_with_synonyms(keyword: str) -> List[str]:
    """
    Expand keyword with known synonyms
    FIX: Ensure we always return at least one variant
    """
    k_norm = normalize_text(keyword)
    if not k_norm:  # Handle empty keywords
        return []
    
    variants = [k_norm]

    for term, syns in TECH_SYNONYMS.items():
        if k_norm == term or k_norm in syns:
            variants.extend([term] + syns)

    # Remove duplicates and empty strings
    return list(set(v for v in variants if v))

def extract_smart_keywords(text: str, top_k: int = 30) -> List[Dict[str, any]]:
    """Extract keywords with importance weights"""
    text_norm = normalize_text(text)

    tokens = re.findall(r"[a-z0-9\+#\.\-]{2,}", text_norm)

    freq: Dict[str, int] = {}
    for t in tokens:
        if t in STOPWORDS or len(t) < 2:
            continue
        freq[t] = freq.get(t, 0) + 1

    scored = []
    for word, count in freq.items():
        weight = 1.0
        if (word in TECH_SYNONYMS) or any(word in syns for syns in TECH_SYNONYMS.values()):
            weight = 2.0
        if any(lang in word for lang in ["python", "java", "javascript", "react", "docker", "kubernetes"]):
            weight = max(weight, 2.5)
        
        # Better scoring: frequency × log(length+1) × weight
        score = count * (1 + math.log1p(len(word))) * weight
        scored.append((word, count, score))

    scored.sort(key=lambda x: x[2], reverse=True)

    keywords = []
    seen = set()
    for word, count, score in scored[:top_k]:
        if word in seen:
            continue
        
        variants = expand_keyword_with_synonyms(word)
        if not variants:  # Skip if no valid variants
            continue
            
        keywords.append({
            "keyword": word,
            "weight": 2.0 if score > 10 else 1.0,
            "variants": variants,
            "count": count
        })
        seen.add(word)

    return keywords

# ------------------------- Experience & Education -------------------------

def _parse_month_year(token: str) -> Tuple[Optional[int], Optional[int]]:
    """Parse month and year from token"""
    t = token.strip().lower()
    if t in ("present", "current", "now", "aujourd'hui"):
        return (None, datetime.now().year)

    # Month Year pattern
    m = re.match(r"([a-zéû\.]+)\s+(\d{4})", t)
    if m:
        mon_raw, year = m.group(1), int(m.group(2))
        mon_raw = mon_raw.strip(".")
        mon = MONTHS.get(mon_raw, None)
        return (mon, year)

    # Year only
    m = re.match(r"(\d{4})", t)
    if m:
        return (None, int(m.group(1)))

    return (None, None)

def _merge_ranges(ranges: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
    """Merge overlapping year ranges"""
    if not ranges:
        return []
    ranges = sorted(ranges, key=lambda r: (r[0], r[1]))
    merged = [ranges[0]]
    for s, e in ranges[1:]:
        last_s, last_e = merged[-1]
        if s <= last_e:
            merged[-1] = (last_s, max(last_e, e))
        else:
            merged.append((s, e))
    return merged

def extract_years_of_experience(text: str) -> Optional[float]:
    """
    Extract total years of experience with improved patterns
    """
    text_norm = normalize_text(text)
    now_year = datetime.now().year

    # 1) Explicit "X years" pattern
    explicit = re.findall(r'(\d+)\+?\s*(?:years?|yrs?)', text_norm)
    explicit = [int(x) for x in explicit if x.isdigit()]
    explicit_years = max(explicit) if explicit else None

    # 2) Date ranges
    ranges = []

    # Pattern: YYYY - YYYY or YYYY - Present
    for m in re.finditer(r'(\d{4})\s*(?:-|to|–|—)\s*(\d{4}|present|current|now)', text_norm):
        s = int(m.group(1))
        end_tok = m.group(2)
        e = now_year if end_tok in ("present", "current", "now") else int(end_tok)
        if 1980 <= s <= now_year and 1980 <= e <= now_year and e >= s:
            ranges.append((s, e))

    # Pattern: Month YYYY - Month YYYY
    for m in re.finditer(r'([A-Za-zéû\.]+)\s+(\d{4})\s*(?:-|to|–|—)\s*([A-Za-zéû\.]+|\d{4}|present|current|now)\s*(\d{4})?', text):
        s_mon_raw, s_year = m.group(1), m.group(2)
        e_mon_raw, e_year_opt = m.group(3), m.group(4)

        _, sy = _parse_month_year(f"{s_mon_raw} {s_year}")
        if e_mon_raw in ("present", "current", "now"):
            ey = now_year
        else:
            if e_year_opt:
                _, ey = _parse_month_year(f"{e_mon_raw} {e_year_opt}")
            else:
                _, ey = _parse_month_year(e_mon_raw)
        
        if sy and ey and 1980 <= sy <= now_year and 1980 <= ey <= now_year and ey >= sy:
            ranges.append((sy, ey))

    # NEW FIX: Pattern "Month YYYY - YYYY" (no month for end)
    for m in re.finditer(r'([A-Za-zéû\.]+)\s+(\d{4})\s*(?:-|to|–|—)\s*(\d{4})', text_norm):
        mon_raw, s_year, e_year = m.group(1), m.group(2), m.group(3)
        sy, ey = int(s_year), int(e_year)
        if 1980 <= sy <= now_year and 1980 <= ey <= now_year and ey >= sy:
            ranges.append((sy, ey))

    total_from_ranges: Optional[float] = None
    if ranges:
        merged = _merge_ranges(ranges)
        # Inclusive years
        total = sum((e - s + 1) for s, e in merged)
        total_from_ranges = float(total)

    # Prefer explicit if present; else use ranges
    if explicit_years is not None:
        return float(explicit_years)
    return total_from_ranges

def detect_seniority_level(text: str, years_exp: Optional[float]) -> str:
    """Detect seniority level"""
    text_norm = normalize_text(text)
    scores = {"entry": 0, "mid": 0, "senior": 0, "expert": 0}

    for level, indicators in SENIORITY_INDICATORS.items():
        for indicator in indicators:
            if indicator in text_norm:
                scores[level] += 1

    if years_exp is not None:
        if years_exp < 2:
            scores["entry"] += 3
        elif years_exp < 5:
            scores["mid"] += 3
        elif years_exp < 10:
            scores["senior"] += 3
        else:
            scores["expert"] += 3

    detected = max(scores.items(), key=lambda x: x[1])
    return detected[0] if detected[1] > 0 else "mid"

def detect_education_level(text: str) -> int:
    """Detect education level (0-100 score)"""
    text_norm = normalize_text(text)
    score = 0
    
    if any(term in text_norm for term in ["phd", "ph.d", "doctorate", "doctoral"]):
        score = 100
    elif any(term in text_norm for term in ["master", "msc", "m.sc", "mba", "m.b.a"]):
        score = 80
    elif any(term in text_norm for term in ["bachelor", "bsc", "b.sc", "ba", "b.a", "bs", "b.s"]):
        score = 60
    elif any(term in text_norm for term in ["associate", "diploma", "college"]):
        score = 40
    elif any(term in text_norm for term in ["high school", "secondary"]):
        score = 20
    
    return score

# ------------------------- Matching & Scoring -------------------------

def _text_chunks(text: str, chunk_len: int) -> List[str]:
    """Split text into chunks"""
    t = text if isinstance(text, str) else str(text)
    if len(t) <= chunk_len:
        return [t]
    
    chunks = []
    i = 0
    while i < len(t):
        chunks.append(t[i:i+chunk_len])
        i += chunk_len
    return chunks

def _make_word_boundary_pattern(variant: str) -> str:
    """
    FIX: Create proper word boundary pattern for special chars like C++, C#, .NET
    Regular \b doesn't work with special characters
    """
    # Check if variant contains only word characters
    if re.match(r'^\w+$', variant):
        # Standard word boundary works
        return rf"\b{re.escape(variant)}\b"
    else:
        # Use custom boundary: not preceded/followed by word chars
        return rf"(?<!\w){re.escape(variant)}(?!\w)"

def fuzzy_match_with_variants(haystack: str, keyword_obj: Dict, threshold: int = 88) -> bool:
    """
    OPTIMIZED: Fuzzy matching with pre-filtering and proper word boundaries
    """
    hay_norm = normalize_text(haystack)

    # FIX: Ensure we have valid variants
    variants = [v for v in (keyword_obj.get("variants", []) or []) if v]
    if not variants:
        # Fallback to keyword itself
        kw = keyword_obj.get("keyword", "")
        if kw:
            variants = [normalize_text(kw)]
        else:
            return False  # Can't match empty keyword

    # Quick exact match with proper word boundaries
    for v in variants:
        pattern = _make_word_boundary_pattern(v)
        if re.search(pattern, hay_norm):
            return True

    # Chunked fuzzy matching with pre-filtering
    for chunk in re.split(r'[.\n]', hay_norm):
        chunk = chunk.strip()
        if not chunk:
            continue
        
        # OPTIMIZATION: Skip very long chunks
        if len(chunk) > FUZZY_MAX_CHUNK_LEN:
            continue
        
        for v in variants:
            # OPTIMIZATION: Pre-filter - check if first 3 chars present
            if FUZZY_PRE_FILTER and len(v) >= 3:
                prefix = v[:3]
                if prefix not in chunk:
                    continue  # Skip expensive fuzzy match
            
            # Now do fuzzy match
            if fuzz.token_set_ratio(chunk, v) >= threshold:
                return True
    
    return False

def compute_keyword_coverage(resume_text: str, jd_keywords: List[Dict]) -> Tuple[float, List[str]]:
    """Compute weighted keyword coverage"""
    matched = []
    total_weight = sum(kw["weight"] for kw in jd_keywords) or 1.0
    matched_weight = 0.0

    for kw_obj in jd_keywords:
        if fuzzy_match_with_variants(resume_text, kw_obj, threshold=88):
            matched.append(kw_obj["keyword"])
            matched_weight += kw_obj["weight"]

    coverage = 100.0 * matched_weight / total_weight
    return coverage, matched

def detect_red_flags(resume_text: str, jd_keywords: List[Dict]) -> List[str]:
    """Detect red flags in resume"""
    flags = []
    text_norm = normalize_text(resume_text)
    words = text_norm.split()

    # FIX: Count JD keywords IN RESUME (not in JD)
    resume_kw_hits = 0
    for kw in jd_keywords:
        k = kw["keyword"]
        # Use proper word boundary
        pattern = _make_word_boundary_pattern(k)
        resume_kw_hits += len(re.findall(pattern, text_norm))

    keyword_density = resume_kw_hits / max(len(words), 1)
    if keyword_density > 0.08:
        flags.append("⚠️ Possible keyword stuffing detected")

    if len(words) < 200:
        flags.append("⚠️ Resume is unusually short (< 200 words)")

    job_count = len(re.findall(r'\b(?:company|employer|organization|inc|corp|llc|sa|sarl|gmbh)\b', text_norm))
    years_exp = extract_years_of_experience(resume_text)
    if years_exp and job_count > years_exp * 2:
        flags.append("⚠️ Frequent job changes detected")

    typo_patterns = [
        r'\b(?:teh|recieve|seperate|occured|managment)\b',
        r'[a-z]{24,}',
    ]
    for pattern in typo_patterns:
        if re.search(pattern, text_norm):
            flags.append("⚠️ Potential spelling/grammar issues")
            break

    return flags

# ------------------------- Engine -------------------------

class ScreeningEngine:
    def __init__(self):
        try:
            import torch
            device = "cuda" if torch.cuda.is_available() else "cpu"
            self.model = SentenceTransformer(MODEL_NAME, device=device)
        except Exception:
            self.model = SentenceTransformer(MODEL_NAME)

    def embed(self, texts: List[str]):
        """Batch embed texts"""
        return self.model.encode(texts, convert_to_tensor=True, normalize_embeddings=True)

    def semantic_score(self, jd_emb, resume_text: str) -> float:
        """
        FIX: Use weighted top-k chunks instead of just max
        """
        if not USE_CHUNKED_SEMANTICS:
            r_emb = self.model.encode(resume_text, convert_to_tensor=True, normalize_embeddings=True)
            return float(util.cos_sim(jd_emb, r_emb).item()) * 100.0

        chunks = _text_chunks(resume_text, SEMANTIC_CHUNK_CHAR_LEN)
        if not chunks:
            return 0.0
        
        c_embs = self.embed(chunks)
        sims = util.cos_sim(jd_emb, c_embs).cpu().numpy().flatten().tolist()
        
        if SEMANTIC_CHUNK_AGG == "avg":
            val = sum(sims) / len(sims)
        elif SEMANTIC_CHUNK_AGG == "max":
            val = max(sims)
        elif SEMANTIC_CHUNK_AGG == "weighted_top":
            # FIX: Weighted average of top-k chunks (more robust)
            sims_sorted = sorted(sims, reverse=True)
            top_k = min(SEMANTIC_TOP_K, len(sims_sorted))
            if top_k == 0:
                val = 0.0
            else:
                val = sum(sims_sorted[:top_k]) / top_k
        else:
            val = max(sims)
        
        return float(val) * 100.0

    def screen_candidate(
        self,
        jd_text: str,
        resume_text: str,
        resume_filename: str,
        must_haves: List[str],
        nice_to_haves: List[str],
        jd_keywords: List[Dict],
        jd_emb,
        required_years: Optional[int] = None,
        required_education: Optional[str] = None,
        weights: ScoringWeights = ScoringWeights()
    ) -> ScreeningResult:

        weights = weights.normalized()

        # Parse must-haves / nice-to-haves
        mh_objects = [{"keyword": normalize_text(m), "weight": 3.0, "variants": expand_keyword_with_synonyms(m)}
                      for m in must_haves if m.strip()]
        # Filter out empty variants
        mh_objects = [obj for obj in mh_objects if obj["variants"]]
        
        nh_objects = [{"keyword": normalize_text(n), "weight": 1.0, "variants": expand_keyword_with_synonyms(n)}
                      for n in nice_to_haves if n.strip()]
        nh_objects = [obj for obj in nh_objects if obj["variants"]]

        # 1) Semantic
        semantic_score = self.semantic_score(jd_emb, resume_text)

        # 2) Keyword coverage
        keyword_coverage, matched_kw = compute_keyword_coverage(resume_text, jd_keywords)

        # 3) Must-haves
        missing_mh = []
        matched_mh_count = 0
        for mh in mh_objects:
            if fuzzy_match_with_variants(resume_text, mh, threshold=88):
                matched_mh_count += 1
            else:
                missing_mh.append(mh["keyword"])
        must_have_score = 100.0 * matched_mh_count / len(mh_objects) if mh_objects else 100.0

        # 4) Experience
        candidate_years = extract_years_of_experience(resume_text)
        _ = detect_seniority_level(resume_text, candidate_years)

        if required_years is not None and required_years < 0:
            required_years = 0

        if required_years:
            if candidate_years is not None:
                if candidate_years >= required_years:
                    experience_score = 100.0
                elif candidate_years >= required_years * 0.7:
                    experience_score = 70.0
                else:
                    experience_score = 30.0
            else:
                experience_score = 50.0
        else:
            experience_score = 100.0

        # 5) Education
        education_score = detect_education_level(resume_text)
        if required_education:
            req_edu_norm = normalize_text(required_education)
            if "bachelor" in req_edu_norm and education_score < 60:
                education_score = max(education_score * 0.5, 20)
            elif "master" in req_edu_norm and education_score < 80:
                education_score = max(education_score * 0.7, 40)
            elif "phd" in req_edu_norm and education_score < 100:
                education_score = max(education_score * 0.8, 60)

        # 6) Nice-to-haves
        matched_nh = []
        for nh in nh_objects:
            if fuzzy_match_with_variants(resume_text, nh, threshold=88):
                matched_nh.append(nh["keyword"])
        bonus_score = min(100.0, len(matched_nh) * 20.0)

        # Red flags
        red_flags = detect_red_flags(resume_text, jd_keywords)

        # Weighted total
        total_score = (
            weights.semantic * semantic_score +
            weights.keywords * keyword_coverage +
            weights.must_haves * must_have_score +
            weights.experience * experience_score +
            weights.education * education_score +
            weights.bonus * bonus_score
        )

        if red_flags:
            total_score *= (1 - 0.05 * len(red_flags))

        decision, confidence = self._make_decision(total_score, must_have_score, semantic_score, red_flags)

        reasoning = self._generate_reasoning(
            decision, total_score, semantic_score, keyword_coverage,
            must_have_score, len(missing_mh), len(matched_nh),
            candidate_years, required_years, red_flags
        )

        strengths = []
        if semantic_score >= 75:
            strengths.append("✅ Strong overall fit with job description")
        if must_have_score >= 90:
            strengths.append("✅ Meets almost all must-have requirements")
        if len(matched_nh) >= 2:
            strengths.append(f"✅ Has {len(matched_nh)} nice-to-have skills")
        if candidate_years and required_years and candidate_years >= required_years * 1.5:
            strengths.append("✅ Exceeds required experience level")

        candidate_name = self._extract_candidate_name(resume_filename)

        return ScreeningResult(
            candidate_name=candidate_name,
            filename=resume_filename,
            decision=decision,
            confidence=round(confidence, 1),
            total_score=round(total_score, 1),
            semantic_score=round(semantic_score, 1),
            keyword_coverage=round(keyword_coverage, 1),
            must_have_score=round(must_have_score, 1),
            experience_score=round(experience_score, 1),
            education_score=round(education_score, 1),
            bonus_score=round(bonus_score, 1),
            years_experience=round(candidate_years, 1) if candidate_years is not None else None,
            matched_keywords=matched_kw[:20],
            missing_must_haves=missing_mh,
            matched_nice_to_haves=matched_nh,
            red_flags=red_flags,
            strengths=strengths,
            reasoning=reasoning
        )

    def _make_decision(self, total_score: float, must_have_score: float, semantic_score: float, red_flags: List[str]) -> Tuple[str, float]:
        if must_have_score < 50 or semantic_score < 40 or len(red_flags) >= 3:
            return "REJECT", 95.0

        if total_score >= 85 and must_have_score >= 80:
            return "STRONG_HIRE", 90.0
        elif total_score >= 75 and must_have_score >= 70:
            return "HIRE", 85.0
        elif total_score >= 60 and must_have_score >= 60:
            return "MAYBE", 70.0
        elif total_score >= 45:
            return "NO_HIRE", 80.0
        else:
            return "REJECT", 90.0

    def _generate_reasoning(
        self, decision: str, total: float, semantic: float, coverage: float,
        must_have: float, missing_mh: int, matched_nh: int,
        years: Optional[float], req_years: Optional[int], flags: List[str]
    ) -> str:
        lines = [f"**Decision: {decision}** (Total Score: {total:.1f}/100)", ""]
        
        if semantic >= 75:
            lines.append(f"✅ Excellent semantic match ({semantic:.1f}%) - resume aligns well with JD")
        elif semantic >= 60:
            lines.append(f"⚠️ Good semantic match ({semantic:.1f}%) but room for improvement")
        else:
            lines.append(f"❌ Weak semantic match ({semantic:.1f}%) - resume doesn't align with JD")

        if coverage >= 70:
            lines.append(f"✅ Strong keyword coverage ({coverage:.1f}%)")
        elif coverage >= 50:
            lines.append(f"⚠️ Moderate keyword coverage ({coverage:.1f}%)")
        else:
            lines.append(f"❌ Low keyword coverage ({coverage:.1f}%)")

        if must_have >= 90:
            lines.append(f"✅ Meets all must-have requirements ({must_have:.1f}%)")
        elif must_have >= 70:
            lines.append(f"⚠️ Missing {missing_mh} must-have requirement(s)")
        else:
            lines.append(f"❌ Missing {missing_mh} critical must-have requirements")

        if req_years and years is not None:
            if years >= req_years:
                lines.append(f"✅ Has {years:.1f} years experience (required: {req_years})")
            else:
                lines.append(f"⚠️ Has {years:.1f} years experience (required: {req_years})")

        if matched_nh > 0:
            lines.append(f"✅ Has {matched_nh} nice-to-have skill(s)")

        if flags:
            lines.append("")
            lines.append("⚠️ **Concerns:**")
            for flag in flags:
                lines.append(f"  {flag}")

        return "\n".join(lines)

    def _extract_candidate_name(self, filename: str) -> str:
        name = os.path.splitext(os.path.basename(filename))[0]
        name = re.sub(r'[_\-]+', ' ', name).strip()
        name = re.sub(r'\s+', ' ', name)
        return name.title() or filename

# ------------------------- FastAPI App -------------------------

app = FastAPI(title="Advanced Resume Screening", version="2.2.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

engine = ScreeningEngine()

# ------------------------- API -------------------------

@app.post("/api/screen")
async def screen_resumes(
    jd_text: Optional[str] = Form(None),
    jd_file: Optional[UploadFile] = File(None),
    must_haves: Optional[str] = Form(None),
    nice_to_haves: Optional[str] = Form(None),
    required_years: Optional[int] = Form(None),
    required_education: Optional[str] = Form(None),
    resumes: List[UploadFile] = File(...),
):
    """Main screening endpoint"""
    if not resumes:
        raise HTTPException(status_code=400, detail="No resumes provided")
    if len(resumes) > MAX_RESUMES:
        raise HTTPException(status_code=400, detail=f"Maximum {MAX_RESUMES} resumes allowed")

    # Get JD text
    if jd_file:
        jd_text_final, _ = extract_text_from_upload(jd_file)
    else:
        jd_text_final = jd_text or ""

    if not jd_text_final.strip():
        raise HTTPException(status_code=400, detail="Job description is empty")

    # Parse lists
    mh_list = [x.strip() for x in re.split(r'[,\n;]+', must_haves or "") if x.strip()]
    nh_list = [x.strip() for x in re.split(r'[,\n;]+', nice_to_haves or "") if x.strip()]

    # Precompute JD features once
    jd_keywords = extract_smart_keywords(jd_text_final, top_k=30)
    jd_emb = engine.embed([jd_text_final])[0:1]

    # Parallel parse resumes
    parsed: List[Tuple[Optional[str], str, Optional[str]]] = []
    with ThreadPoolExecutor(max_workers=MAX_PARSE_WORKERS) as ex:
        futures = {ex.submit(extract_text_from_upload, rf): rf for rf in resumes}
        for fut in as_completed(futures):
            rf = futures[fut]
            try:
                text, name = fut.result()
                parsed.append((text, name, None))
            except HTTPException as e:
                parsed.append((None, rf.filename or "file", f"Error processing file: {e.detail}"))
            except Exception as e:
                parsed.append((None, rf.filename or "file", f"Unexpected error: {str(e)}"))

    # Screen each resume
    results_json = []
    for text, name, err in parsed:
        if err:
            results_json.append({
                "candidate_name": name,
                "filename": name,
                "decision": "ERROR",
                "confidence": 0,
                "reasoning": err
            })
            continue

        try:
            result = engine.screen_candidate(
                jd_text=jd_text_final,
                resume_text=text,
                resume_filename=name,
                must_haves=mh_list,
                nice_to_haves=nh_list,
                jd_keywords=jd_keywords,
                jd_emb=jd_emb,
                required_years=required_years,
                required_education=required_education
            )
            results_json.append(result.to_dict())
        except Exception as e:
            results_json.append({
                "candidate_name": name,
                "filename": name,
                "decision": "ERROR",
                "confidence": 0,
                "reasoning": f"Error during screening: {str(e)}"
            })

    # Sort by total score
    results_json.sort(key=lambda x: x.get("total_score", -1), reverse=True)

    strong_hire = [r for r in results_json if r.get("decision") == "STRONG_HIRE"]
    hire = [r for r in results_json if r.get("decision") == "HIRE"]
    maybe = [r for r in results_json if r.get("decision") == "MAYBE"]
    no_hire = [r for r in results_json if r.get("decision") == "NO_HIRE"]
    reject = [r for r in results_json if r.get("decision") == "REJECT"]

    return {
        "summary": {
            "total_candidates": len(results_json),
            "strong_hire": len(strong_hire),
            "hire": len(hire),
            "maybe": len(maybe),
            "no_hire": len(no_hire),
            "reject": len(reject),
        },
        "results": results_json,
        "grouped": {
            "strong_hire": strong_hire,
            "hire": hire,
            "maybe": maybe,
            "no_hire": no_hire,
            "reject": reject,
        }
    }

@app.get("/api/health")
def health():
    return {
        "status": "healthy",
        "model": MODEL_NAME,
        "timestamp": time.time(),
        "version": "2.2.0"
    }
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import os

# Serve static files
STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")

# Mount static assets (CSS, JS, images)
if os.path.exists(STATIC_DIR):
    app.mount("/assets", StaticFiles(directory=os.path.join(STATIC_DIR, "assets")), name="assets")

# Serve index.html at root
@app.get("/", include_in_schema=False)
def root():
    index_path = os.path.join(STATIC_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"error": "index.html not found. Please create static/index.html"}

# Serve guide
@app.get("/guide", include_in_schema=False)
def guide():
    guide_path = os.path.join(STATIC_DIR, "important.html")
    if os.path.exists(guide_path):
        return FileResponse(guide_path)
    return {"error": "important.html not found"}

# ========== END OF ADDED CODE ==========

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)